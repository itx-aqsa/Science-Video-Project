from dotenv import load_dotenv
import os
import wikipediaapi # type: ignore
from groq import Groq
import requests
from gtts import gTTS
import PyPDF2
from docx import Document
from langdetect import detect
import tempfile
import wikipedia
import io

load_dotenv()

# Set API keys
Groq_API = os.getenv("Groq_API_Key")
client = Groq(api_key=Groq_API)

GOOGLE_TRANSLATE_API_KEY = os.getenv("Google_API")


def fetch_wikipedia_summary(topic):
    wiki_wiki = wikipediaapi.Wikipedia(user_agent="EducationalScriptApp/1.0", language="en")
    page = wiki_wiki.page(topic)

    if page.exists():
        return page.summary

    try:
        search_results = wikipedia.search(topic, results=3)
        for related_topic in search_results:
            page = wiki_wiki.page(related_topic)
            if page.exists():
                return page.summary
        return None
    except:
        return None


def generate_script(topic, duration):
    factual_content = fetch_wikipedia_summary(topic)
    words_per_minute = 130
    target_words = duration * words_per_minute
    if factual_content:
         # Wikipedia content available, format it into a script
         prompt = f"Format the following factual content into an educational script in English with approximately {target_words} words. Do not include extra text like 'Here is the formatted script' or descriptions.\n\n{factual_content}"
    else:
        # No Wikipedia content, generate a script directly
        prompt = f"Generate an educational script in English on the topic '{topic}' with approximately {target_words} words. Ensure the content is factual, engaging, and suitable for an educational video. Do not include extra text like 'Here is the formatted script' or descriptions."

    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": "You are an AI assistant that creates structured educational scripts. Start directly with the script content, focusing on the topic. Do NOT include introductory phrases, word counts, or metadata."},
            {"role": "user", "content": prompt}
        ],
        model="llama3-70b-8192"
    )

    script = response.choices[0].message.content.strip()
    script = script.replace("**", "").replace("*", "").replace("###", "").replace("##", "").replace("#", "")
    return script       
        

def generate_video_script(script_content):
    response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": "You are an AI assistant that converts educational text into video scene descriptions."},
            {"role": "user", "content": f"Convert this script into a video script with scene descriptions:\n\n{script_content}"}
        ],
        model="llama3-70b-8192"
    )
    return response.choices[0].message.content.strip()


def translate_script(script_content, target_language):
    """
    Translate script using multiple methods - Google API with fallback to Groq AI
    """
    if not script_content:
        return "⚠️ No valid script content to translate."

    # Language mapping for Google Translate API
    language_map = {
        "chinese-simplified": "zh-cn",
        "chinese-traditional": "zh-tw", 
        "spanish": "es",
        "french": "fr",
        "german": "de",
        "italian": "it",
        "japanese": "ja",
        "turkish": "tr",
        "russian": "ru",
        "urdu": "ur",  # Added Urdu
        "arabic": "ar",
        "hindi": "hi",
        "portuguese": "pt",
        "korean": "ko"
    }
    
    google_lang_code = language_map.get(target_language, target_language)
    
    # Skip Google API for now and go directly to Groq AI (more reliable)
    print(f"🤖 Using Groq AI translation for {target_language}")
    
    try:
        return translate_with_groq_ai(script_content, target_language)
    except Exception as e:
        print(f"❌ Groq AI translation failed: {str(e)}")
        return f"❌ Translation failed for {target_language}. Please try again later."


def translate_with_groq_ai(script_content, target_language):
    """
    Translate using Groq AI as primary method
    """
    # Language mapping for human-readable names
    language_names = {
        "chinese-simplified": "Chinese (Simplified)",
        "chinese-traditional": "Chinese (Traditional)",
        "spanish": "Spanish",
        "french": "French", 
        "german": "German",
        "italian": "Italian",
        "japanese": "Japanese",
        "turkish": "Turkish",
        "russian": "Russian",
        "urdu": "Urdu",  # Added Urdu
        "arabic": "Arabic",
        "hindi": "Hindi",
        "portuguese": "Portuguese",
        "korean": "Korean"
    }
    
    target_lang_name = language_names.get(target_language, target_language.title())
    
    # Split long content into chunks for better translation
    chunks = split_text_into_chunks(script_content, max_length=2000)
    translated_chunks = []
    
    for i, chunk in enumerate(chunks):
        print(f"🔄 Translating chunk {i+1}/{len(chunks)} to {target_lang_name}")
        
        prompt = f"""Translate the following educational script from English to {target_lang_name}. 
        Maintain the educational tone, structure, and formatting. Provide only the translation without any additional text or explanations.

        Text to translate:
        {chunk}"""

        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": f"You are a professional translator specializing in educational content. Translate accurately to {target_lang_name} while maintaining the original structure and educational tone. Provide only the translation."},
                {"role": "user", "content": prompt}
            ],
            model="llama3-70b-8192",
            temperature=0.3  # Lower temperature for more consistent translation
        )

        translated_chunk = response.choices[0].message.content.strip()
        translated_chunks.append(translated_chunk)
    
    # Join all translated chunks
    full_translation = "\n\n".join(translated_chunks)
    
    # Add a note that this was AI translated
    full_translation += f"\n\n---\n*Translated to {target_lang_name} using AI*"
    
    return full_translation


def split_text_into_chunks(text, max_length=2000):
    """
    Split text into smaller chunks for better translation
    """
    if len(text) <= max_length:
        return [text]
    
    chunks = []
    paragraphs = text.split('\n\n')
    current_chunk = ""
    
    for paragraph in paragraphs:
        if len(current_chunk + paragraph) <= max_length:
            current_chunk += paragraph + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = paragraph + "\n\n"
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks


def extract_text_from_file(file_content, filename, content_type):
    """
    Extract text from different file types using file content and metadata
    """
    try:
        # Get file extension from filename
        if filename:
            extension = os.path.splitext(filename)[1].lower()
        else:
            # Fallback to content type
            if content_type == "text/plain":
                extension = ".txt"
            elif content_type == "application/pdf":
                extension = ".pdf"
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                extension = ".docx"
            else:
                extension = ".txt"  # Default fallback
        
        print(f"📄 Processing file with extension: {extension}")
        
        if extension == ".txt":
            # Handle text files
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    return file_content.decode("latin-1")
                except UnicodeDecodeError:
                    return file_content.decode("utf-8", errors="ignore")
        
        elif extension == ".pdf":
            # Handle PDF files
            try:
                pdf_file = io.BytesIO(file_content)
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as e:
                print(f"⚠️ PDF extraction error: {str(e)}")
                return f"Error extracting PDF content: {str(e)}"
        
        elif extension == ".docx":
            # Handle DOCX files
            try:
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text.strip()
            except Exception as e:
                print(f"⚠️ DOCX extraction error: {str(e)}")
                return f"Error extracting DOCX content: {str(e)}"
        
        else:
            # Fallback for unknown file types
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                return file_content.decode("utf-8", errors="ignore")
                
    except Exception as e:
        print(f"❌ Text extraction error: {str(e)}")
        return f"Error extracting text: {str(e)}"


def text_to_speech_from_content(text_content):
    """
    Convert text content directly to speech and return audio data
    """
    try:
        if not text_content.strip():
            return "No text found.", None
            
        # Detect language
        try:
            language = detect(text_content)
            print(f"🔍 Detected language: {language}")
        except:
            language = "en"  # Default to English
            print("⚠️ Language detection failed, using English")
        
        # Generate TTS
        tts = gTTS(text=text_content, lang=language, slow=False)
        
        # Save to BytesIO instead of file
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)
        audio_data = audio_buffer.read()
        
        print(f"✅ TTS generated: {len(audio_data)} bytes")
        
        return f"Language detected: {language}", audio_data
        
    except Exception as e:
        print(f"❌ TTS error: {str(e)}")
        return f"TTS Error: {str(e)}", None


def extract_text(file):
    """
    Legacy function for backward compatibility
    """
    try:
        if hasattr(file, 'name') and file.name:
            extension = os.path.splitext(file.name)[1].lower()
        else:
            # Try to read as text file
            extension = ".txt"
            
        if extension == ".txt":
            content = file.read()
            if isinstance(content, bytes):
                return content.decode("utf-8")
            return content
        elif extension == ".pdf":
            reader = PyPDF2.PdfReader(file)
            return "".join(page.extract_text() for page in reader.pages)
        elif extension == ".docx":
            doc = Document(file)
            return "\n".join(para.text for para in doc.paragraphs)
        else:
            # Try to read as text
            content = file.read()
            if isinstance(content, bytes):
                return content.decode("utf-8")
            return content
    except Exception as e:
        print(f"❌ Legacy extract_text error: {str(e)}")
        raise ValueError(f"Error extracting text from file: {str(e)}")


def text_to_speech(file):
    """
    Legacy function - convert file to speech
    """
    try:
        text = extract_text(file)
        return text_to_speech_from_content(text)
    except Exception as e:
        print(f"❌ Legacy TTS error: {str(e)}")
        return f"Error: {str(e)}", None
